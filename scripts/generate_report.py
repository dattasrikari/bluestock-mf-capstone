"""
Generates the final capstone report (Final_Report.docx) for Bluestock Fintech
Mutual Fund Analytics Platform, pulling in charts, dashboard screenshots,
and data tables from across the project.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import pandas as pd

BASE_DIR = Path(__file__).resolve().parent.parent
CHARTS_DIR = BASE_DIR / "charts"
DASHBOARD_DIR = BASE_DIR / "dashboard"
REPORTS_DIR = BASE_DIR / "reports"
OUTPUT_PATH = BASE_DIR / "reports" / "Final_Report.docx"

PURPLE = RGBColor(0x6B, 0x2C, 0x91)

doc = Document()

# ---- Title Page ----
title = doc.add_heading("Bluestock Fintech", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = PURPLE

subtitle = doc.add_paragraph("Mutual Fund Analytics Platform")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(20)
subtitle.runs[0].font.bold = True

sub2 = doc.add_paragraph("Capstone Project - Final Report")
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.runs[0].font.size = Pt(14)

doc.add_paragraph()
meta = doc.add_paragraph("Prepared by: Intern / Data Analyst - Bluestock Fintech\nDate: August 2026")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

print("Title page created.")

# ---- Table of Contents ----
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Problem Statement",
    "3. Data Sources and Datasets",
    "4. System Architecture and ETL Pipeline",
    "5. Exploratory Data Analysis Findings",
    "6. Performance Analysis",
    "7. Dashboard Overview",
    "8. Advanced Analytics and Risk Metrics",
    "9. Recommendations",
    "10. Limitations",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ---- 1. Executive Summary ----
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This report presents the Bluestock Fintech Mutual Fund Analytics Platform, a full-stack "
    "data engineering and analytics project built over a 7-day capstone. The project ingests "
    "publicly available data from AMFI India, mfapi.in, and NSE/BSE sources, transforms it "
    "through a Python-based ETL pipeline, stores it in a normalized SQLite database, and "
    "presents insights through an interactive Power BI dashboard."
)
doc.add_paragraph(
    "The platform covers 40 real mutual fund schemes across major Indian fund houses, with "
    "over 64,000 NAV records spanning January 2022 to May 2026, and simulated investor "
    "transaction data for 5,000 investors. Key analytical outputs include fund performance "
    "metrics (Sharpe, Sortino, Alpha, Beta, Maximum Drawdown), risk metrics (Value at Risk, "
    "Conditional VaR, sector concentration via Herfindahl-Hirschman Index), investor behavior "
    "analysis (cohort analysis, SIP continuity), and a simple risk-based fund recommendation engine."
)
doc.add_paragraph(
    "The dashboard delivers four interactive report pages (Industry Overview, Fund Performance, "
    "Investor Analytics, and SIP & Market Trends) plus a drill-in NAV detail view, all built "
    "with Bluestock Fintech branding and cross-filtering slicers."
)

doc.add_page_break()
print("Table of Contents and Executive Summary added.")

# ---- 2. Problem Statement ----
doc.add_heading("2. Problem Statement", level=1)
doc.add_paragraph(
    "Despite the rapid growth of India's mutual fund industry, individual investors and "
    "financial advisors face several challenges in making data-driven fund selection decisions:"
)

problems = [
    ("Data Fragmentation", "NAV, AUM, SIP flow, and portfolio holding data are published across "
     "different AMFI sources in different formats, with no unified database."),
    ("Performance Comparison Gap", "Investors cannot easily compare funds across AMCs on a "
     "risk-adjusted basis without significant manual computation of Sharpe ratio, Alpha, and Beta."),
    ("No Benchmark Tracking", "Most retail investors lack visibility into whether their fund is "
     "outperforming its benchmark index."),
    ("Investor Behaviour Blind Spot", "AMCs have limited visibility into how demographic and "
     "geographic factors influence SIP amounts and redemption patterns."),
    ("Slow Reporting", "Monthly MF reports are static PDFs; stakeholders need live, self-service "
     "dashboards with drill-down capability."),
]
for name, desc in problems:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "This platform addresses each of these gaps through an automated ETL pipeline, a normalized "
    "analytical database, computed risk-adjusted performance metrics, and an interactive "
    "self-service dashboard."
)

doc.add_page_break()

# ---- 3. Data Sources and Datasets ----
doc.add_heading("3. Data Sources and Datasets", level=1)
doc.add_paragraph(
    "All data used in this project is derived from publicly available sources, primarily AMFI "
    "India, mfapi.in, and NSE/BSE public reports. The table below summarizes the core datasets "
    "used across the pipeline."
)

dataset_table_data = [
    ("Dataset", "Rows", "Description"),
    ("01_fund_master.csv", "40", "Master list of mutual fund schemes with AMFI codes, fund house, category"),
    ("02_nav_history.csv", "~64,320", "Daily NAV for all 40 schemes, Jan 2022 - May 2026"),
    ("03_aum_by_fund_house.csv", "~90", "Quarterly AUM by fund house"),
    ("04_monthly_sip_inflows.csv", "48", "Monthly SIP inflow, active accounts, AUM"),
    ("05_category_inflows.csv", "~144", "Net inflows by fund category"),
    ("06_industry_folio_count.csv", "21", "Total MF folios by category"),
    ("08_investor_transactions.csv", "~32,778", "Simulated investor transactions with demographics"),
    ("09_portfolio_holdings.csv", "~322", "Top equity holdings per fund"),
    ("10_benchmark_indices.csv", "~8,000", "Daily benchmark index values (Nifty 50/100, etc.)"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
for i, header in enumerate(dataset_table_data[0]):
    hdr_cells[i].text = header
    hdr_cells[i].paragraphs[0].runs[0].bold = True

for row_data in dataset_table_data[1:]:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

doc.add_page_break()
print("Problem Statement and Data Sources added.")

# ---- 4. System Architecture and ETL Pipeline ----
doc.add_heading("4. System Architecture and ETL Pipeline", level=1)
doc.add_paragraph(
    "The platform follows a standard data engineering architecture: Extract, Transform, Load, "
    "Analyse, and Visualise. This mirrors real-world fintech data pipelines."
)

layers = [
    ("Layer 1 - Data Sources (Extract)", "AMFI daily NAV files, mfapi.in REST API (live NAV "
     "fetched for 6 real schemes including HDFC Top 100 and SBI Bluechip), and 10 pre-packaged "
     "CSV datasets."),
    ("Layer 2 - Data Processing (Transform)", "Python (Pandas) scripts parse, clean, and reshape "
     "all datasets: forward-filling missing NAV values on non-trading days, computing daily "
     "returns, normalising fund names, and validating AMFI codes against the master list."),
    ("Layer 3 - Data Storage (Load)", "A 5-table SQLite star schema (dim_fund, fact_nav, "
     "fact_transactions, fact_performance, fact_aum) stores the cleaned data, indexed for fast "
     "query performance."),
    ("Layer 4 - Analytics (Analyse)", "Jupyter notebooks compute performance metrics (Sharpe, "
     "Sortino, Alpha, Beta, Maximum Drawdown), risk metrics (VaR, CVaR, sector HHI), and investor "
     "behaviour analysis (cohort analysis, SIP continuity)."),
    ("Layer 5 - Visualisation (Dashboard)", "A Power BI dashboard with 4 report pages plus a "
     "drill-in NAV detail view, branded with Bluestock Fintech's identity and cross-filtering "
     "slicers."),
]
for name, desc in layers:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_heading("Database Schema", level=2)
schema_data = [
    ("Table", "Type", "Rows"),
    ("dim_fund", "Dimension", "40"),
    ("fact_nav", "Fact", "~64,320"),
    ("fact_transactions", "Fact", "~32,778"),
    ("fact_performance", "Fact", "40"),
    ("fact_aum", "Fact", "~90"),
]
schema_table = doc.add_table(rows=1, cols=3)
schema_table.style = "Light Grid Accent 1"
schema_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = schema_table.rows[0].cells
for i, h in enumerate(schema_data[0]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for row in schema_data[1:]:
    cells = schema_table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_page_break()
print("System Architecture section added.")

# ---- 5. Exploratory Data Analysis Findings ----
doc.add_heading("5. Exploratory Data Analysis Findings", level=1)
doc.add_paragraph(
    "Exploratory analysis was performed on NAV, AUM, SIP, and investor transaction data to "
    "identify key trends and patterns ahead of deeper performance and risk analysis."
)

def add_chart(filename, caption, width_inches=6.0):
    path = CHARTS_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Chart not found: {filename}]")

add_chart("nav_trend_all_schemes.png", "Figure 1: NAV trend for all 40 schemes, 2022-2026")
add_chart("aum_growth_by_amc.png", "Figure 2: AUM growth by fund house, 2022-2025")
add_chart("sip_inflow_trend.png", "Figure 3: Monthly SIP inflow trend")
doc.add_page_break()

add_chart("category_inflow_heatmap.png", "Figure 4: Category-wise inflow heatmap")
add_chart("demographics.png", "Figure 5: Investor demographics - age group and SIP amount distribution")
doc.add_page_break()

add_chart("geographic_distribution.png", "Figure 6: Geographic distribution of SIP amounts by state")
add_chart("folio_growth.png", "Figure 7: Total MF folio count growth, 2022-2025")
doc.add_page_break()

add_chart("correlation_matrix.png", "Figure 8: Correlation matrix of NAV returns across selected funds")
add_chart("sector_allocation_donut.png", "Figure 9: Sector allocation across equity fund portfolios")
doc.add_page_break()

doc.add_heading("Key EDA Findings", level=2)
eda_findings = [
    "SBI Mutual Fund dominates AUM, roughly 40% higher than the second-placed ICICI Prudential MF.",
    "Industry-wide SIP inflows peaked at an all-time high of Rs. 31,002 crore in December 2025.",
    "Liquid funds dominate category-wise net inflows relative to equity categories.",
    "The 26-35 age group makes up 41.1% of investors, but the 56+ age group shows the highest "
    "median SIP amount, suggesting older investors commit larger sums per SIP.",
    "Total MF folios nearly doubled over the observed period, from 13.26 crore to 26.12 crore.",
    "Correlations between fund NAV daily returns are close to zero across most fund pairs, "
    "indicating the simulated NAV series behave largely independently of one another.",
    "Banking, IT, and Pharma sectors together account for approximately 45% of combined equity "
    "portfolio holdings across all funds.",
]
for finding in eda_findings:
    doc.add_paragraph(finding, style="List Bullet")

doc.add_page_break()
print("EDA Findings section added.")

# ---- 6. Performance Analysis ----
doc.add_heading("6. Performance Analysis", level=1)
doc.add_paragraph(
    "Fund-level performance and risk-adjusted return metrics were computed from the cleaned "
    "NAV history, including CAGR, Sharpe ratio, Sortino ratio, Alpha, Beta, and Maximum "
    "Drawdown. A composite Fund Scorecard (0-100) was then built to rank all 40 schemes."
)

doc.add_heading("Top 10 Funds by Scorecard", level=2)
scorecard = pd.read_csv(REPORTS_DIR / "fund_scorecard.csv")
top10_scorecard = scorecard.sort_values("fund_score", ascending=False).head(10)

sc_table = doc.add_table(rows=1, cols=5)
sc_table.style = "Light Grid Accent 1"
sc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Scheme Name", "Category", "Sharpe", "3Yr CAGR", "Fund Score"]
hdr = sc_table.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True

for _, row in top10_scorecard.iterrows():
    cells = sc_table.add_row().cells
    cells[0].text = str(row["scheme_name"])[:40]
    cells[1].text = str(row["category"])
    cells[2].text = f"{row['sharpe_ratio']:.3f}"
    cells[3].text = f"{row['cagr_3yr']*100:.2f}%"
    cells[4].text = f"{row['fund_score']:.2f}"

doc.add_paragraph()
doc.add_paragraph(
    "Mirae Asset Large Cap Fund ranks highest overall with a fund score of 86.25, driven by "
    "strong Sharpe ratio and 3-year CAGR relative to its peers."
)

doc.add_page_break()

add_chart("benchmark_comparison_top5.png", "Figure 10: Top 5 funds vs. benchmark index, 3-year comparison")
add_chart("daily_returns_distribution.png", "Figure 11: Distribution of daily returns across funds")

doc.add_paragraph()
doc.add_heading("Key Performance Findings", level=2)
perf_findings = [
    "Alpha and Beta were computed against NIFTY100 via OLS regression; r-squared values were "
    "close to zero across most funds, indicating minimal linear relationship between fund "
    "returns and the benchmark in this dataset.",
    "Small Cap and Mid Cap equity funds show the deepest maximum drawdowns, with SBI Small "
    "Cap Fund recording the worst at approximately -52.6%.",
    "Tracking error for the top 5 scorecard funds ranges from 19% to 23% relative to their "
    "benchmark indices.",
]
for f in perf_findings:
    doc.add_paragraph(f, style="List Bullet")

doc.add_page_break()
print("Performance Analysis section added.")

# ---- 7. Dashboard Overview ----
doc.add_heading("7. Dashboard Overview", level=1)
doc.add_paragraph(
    "An interactive Power BI dashboard was built to give stakeholders a self-service view of "
    "the mutual fund industry, individual fund performance, investor behaviour, and SIP/market "
    "trends. The dashboard consists of five pages, described below."
)

def add_dashboard_page(filename, caption, width_inches=6.3):
    path = DASHBOARD_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Dashboard page not found: {filename}]")

doc.add_heading("Page 1: Industry Overview", level=2)
doc.add_paragraph(
    "KPI cards for Total Folios, Total AUM, Total SIP Inflows, and Total Schemes, alongside an "
    "AUM growth trend line and AUM by fund house bar chart."
)
add_dashboard_page("Page1_IndustryOverview.png", "Figure 12: Industry Overview dashboard page")
doc.add_page_break()

doc.add_heading("Page 2: Fund Performance", level=2)
doc.add_paragraph(
    "A sortable Fund Scorecard table, a Risk vs Return scatter plot sized by expense ratio, "
    "and a NAV vs Benchmark comparison chart, filterable by fund house, category, and plan."
)
add_dashboard_page("Page2_FundPerformance.png", "Figure 13: Fund Performance dashboard page")
doc.add_page_break()

doc.add_heading("Page 3: Investor Analytics", level=2)
doc.add_paragraph(
    "Transaction amount by state, transaction type split (SIP/Lumpsum/Redemption), average SIP "
    "amount by age group, and monthly transaction volume, filterable by state, age group, and "
    "city tier."
)
add_dashboard_page("Page3_InvestorAnalytics.png", "Figure 14: Investor Analytics dashboard page")
doc.add_page_break()

doc.add_heading("Page 4: SIP & Market Trends", level=2)
doc.add_paragraph(
    "A dual-axis chart comparing SIP inflow against Nifty 50 index movement, a category-wise "
    "inflow heatmap, and the top 5 fund categories by net inflow for FY 2024-25."
)
add_dashboard_page("Page4_SIPMarketTrends.png", "Figure 15: SIP & Market Trends dashboard page")
doc.add_page_break()

doc.add_heading("Page 5: NAV Detail (Interactive Drill-in)", level=2)
doc.add_paragraph(
    "A fund-level NAV history view, accessed via a 'View NAV Detail' navigation button from "
    "the Fund Performance page, with a scheme selector to inspect any individual fund's NAV "
    "trend."
)
add_dashboard_page("Page5_NAVDetails.png", "Figure 16: NAV Detail dashboard page")
doc.add_page_break()

print("Dashboard Overview section added.")

# ---- 8. Advanced Analytics and Risk Metrics ----
doc.add_heading("8. Advanced Analytics and Risk Metrics", level=1)
doc.add_paragraph(
    "Building on the core performance metrics, additional risk and behavioural analyses were "
    "performed to give a deeper view of downside risk, investor patterns, and portfolio "
    "concentration."
)

doc.add_heading("Value at Risk and Conditional VaR", level=2)
doc.add_paragraph(
    "Historical VaR (95%) and CVaR were computed for each fund from its daily return "
    "distribution. Small Cap and Mid Cap equity funds carry the highest downside risk, with "
    "ABSL Small Cap Fund showing the worst 95% VaR at -2.39% daily and CVaR of -3.03%, "
    "consistent with the deep maximum drawdowns observed in the performance analysis."
)

add_chart("rolling_sharpe_chart.png", "Figure 17: Rolling 90-day Sharpe ratio for 5 selected funds")
doc.add_paragraph(
    "The rolling Sharpe analysis reveals that Liquid funds show persistently negative Sharpe "
    "ratios throughout the observed period, as their low absolute returns fail to clear the "
    "6.5% risk-free benchmark, while Small Cap funds oscillate around and above zero despite "
    "higher volatility."
)

doc.add_page_break()

doc.add_heading("Investor Cohort Analysis", level=2)
doc.add_paragraph(
    "Investors were grouped by the year of their first transaction. The dataset captures two "
    "cohort years: 4,803 investors starting in 2024 and 197 in 2025 (transaction data spans "
    "approximately January 2024 to May 2025). Both cohorts show a preference for Equity "
    "category funds. Notably, the smaller 2025 cohort shows a higher average SIP amount "
    "(Rs. 13,505) than the 2024 cohort (Rs. 10,997), suggesting newer investors may be "
    "starting with stronger initial commitments."
)

doc.add_heading("SIP Continuity Analysis", level=2)
doc.add_paragraph(
    "For investors with 6 or more SIP transactions, the average gap between transactions was "
    "computed to flag continuity risk. The brief's fixed 35-day threshold proved unsuitable for "
    "this dataset - the median SIP gap is 64.7 days, well above 35 days, meaning a literal "
    "application flags 97.8% of eligible investors, which is not a meaningful signal. "
    "Recalibrating to a relative threshold (the top 25% of investors by gap width, greater than "
    "75.6 days) produces a more actionable at-risk figure of 25% (341 of 1,362 eligible "
    "investors)."
)

doc.add_heading("Fund Recommendation Engine", level=2)
doc.add_paragraph(
    "A simple recommendation function (scripts/recommender.py) matches an investor's stated "
    "risk appetite (Low, Moderate, or High) to the dataset's actual risk categories and returns "
    "the top 3 funds by Sharpe ratio within that band. Notably, all three top 'Low' risk funds "
    "(Gilt and Short Duration debt funds) show negative Sharpe ratios, consistent with the "
    "rolling Sharpe finding - low-risk debt instruments in this dataset did not generate enough "
    "return to clear the risk-free rate."
)

doc.add_page_break()

doc.add_heading("Sector Concentration (HHI)", level=2)
add_chart("sector_hhi_chart.png", "Figure 18: Top 10 most sector-concentrated equity funds")
doc.add_paragraph(
    "The Herfindahl-Hirschman Index was computed per fund from its sector-level portfolio "
    "weights. Axis Bluechip Fund is the most sector-concentrated (HHI of 0.297), and even the "
    "Nifty 50 index ETF shows notable concentration (0.237), a natural consequence of "
    "large-cap and index funds weighting toward the same dominant sectors."
)

doc.add_page_break()
print("Advanced Analytics section added.")

# ---- 9. Recommendations ----
doc.add_heading("9. Recommendations", level=1)
doc.add_paragraph(
    "Based on the analysis performed across this platform, the following recommendations are "
    "made for Bluestock Fintech and its stakeholders:"
)

recommendations = [
    "Prioritise large-cap and multi-cap funds for risk-averse investors. Funds such as Mirae "
    "Asset Large Cap (fund score 86.25) offer the strongest risk-adjusted returns in this "
    "dataset, making them suitable anchor holdings for conservative portfolios.",
    "Re-evaluate low-risk debt fund allocations. Gilt and Short Duration debt funds in this "
    "dataset consistently underperformed the risk-free rate, which is atypical - if this "
    "pattern holds in production data, it warrants closer review of fund selection within the "
    "'Low' risk category.",
    "Focus SIP retention efforts on the top 25% widest-gap investors. Rather than applying a "
    "blanket 35-day rule, a data-driven relative threshold more accurately identifies the "
    "investors most likely to lapse.",
    "Monitor sector concentration in flagship large-cap funds. Axis Bluechip Fund and similar "
    "index-tracking products carry meaningfully higher sector concentration risk than actively "
    "diversified peers, which should be disclosed clearly to investors.",
    "Expand the recommendation engine with additional inputs (investment horizon, existing "
    "portfolio composition) beyond risk appetite alone, to improve personalisation.",
]
for r in recommendations:
    doc.add_paragraph(r, style="List Bullet")

doc.add_page_break()

# ---- 10. Limitations ----
doc.add_heading("10. Limitations", level=1)
doc.add_paragraph(
    "This project has several limitations that should be considered when interpreting its "
    "findings:"
)

limitations = [
    "Investor transaction data is simulated, not real. While the underlying demographic and "
    "geographic distributions were built to reflect realistic patterns, the specific "
    "transaction-level behaviour (e.g., SIP continuity gaps) does not represent actual investor "
    "behaviour and should not be used for production decision-making without validation "
    "against real data.",
    "The transaction dataset spans a short window (approximately January 2024 to May 2025), "
    "limiting the reliability of year-over-year cohort comparisons - only two cohort years "
    "exist in the data.",
    "NAV history, while anchored to real AMFI values at select points, is partly simulated "
    "forward using modelled return and volatility parameters, so fund-to-fund correlations "
    "and benchmark relationships (Alpha, Beta, r-squared) may not match real market behaviour.",
    "The fund recommendation engine uses a single metric (Sharpe ratio) within a risk band; a "
    "production system would benefit from a multi-factor scoring approach and investor-specific "
    "constraints such as investment horizon and liquidity needs.",
    "The dashboard's native drill-through feature was found to be unreliable within Power BI "
    "Desktop for this specific file; a button-based navigation workaround was implemented "
    "instead, which is fully functional but differs from the standard drill-through pattern.",
]
for l in limitations:
    doc.add_paragraph(l, style="List Bullet")

doc.add_paragraph()
doc.add_paragraph(
    "This project is intended for educational purposes and skill development, and does not "
    "constitute financial advice. Mutual fund investments are subject to market risks."
).runs[0].italic = True

print("Recommendations and Limitations sections added.")

# ---- SAVE (keep this at the very bottom as we add more sections) ----
doc.save(OUTPUT_PATH)
print(f"Saved draft to {OUTPUT_PATH}")